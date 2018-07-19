# coding=utf-8
from monitor.util.mysql_util import getconn, closeAll, md5
from monitor import app
import os,sys,time
from docx import Document
from monitor.main.measure.langconv import Converter
from monitor.util.config import upload_files
ALLOWED_EXTENSIONS = set(['png', 'jpg', 'jpeg', 'gif', 'docx'])


# 校验文件是否正确
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS


# 繁体转简体函数
def cht_to_chs(line):
    line = Converter('zh-hans').convert(line)
    line.encode('utf-8')
    return line


# 覆盖掉原来的繁体文件并插入数据库
def save_new_file(file_root, file_name, user_name, is_private, file_label):
    conn = getconn()
    cur = conn.cursor()

    file_root = file_root
    file_name = file_name
    user_name = user_name
    is_private = is_private
    insert_sql = """INSERT INTO `word_files` (`user`,`word_name`,`is_private`,`file_label`) VALUES (%s,%s,%s,%s)"""
    select_sql = """SELECT `id` FROM `word_files` WHERE `word_name` = %s """

    try:
        select_re = cur.execute(select_sql, (file_name))
        if select_re != 0:
            return 2
        insert_re = cur.execute(insert_sql, (user_name, file_name, is_private, file_label))
        if insert_re < 1:
            return 0
        else:
            document = Document(r'{}'.format(file_root))
            document_new = Document()
            for paragraph in document.paragraphs:
                line = paragraph.text
                new_line = cht_to_chs(line)
                document_new.add_paragraph('{}'.format(new_line))
            document_new.save(r'{}'.format(file_root))
            return 1
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


# 删除库中文件
def process_delete(file_name, user_name):
    # 判断是否是docx结尾，避免删掉系统py文件
    conn = getconn()
    cur = conn.cursor()
    word_name = file_name
    user = user_name
    try:
        file_name_list = file_name.split('.')
        if file_name_list[1] != "docx":
            app.logger.error("传入文件格式不对")
            return 0
        # 查询该文件是否属于该用户
        select_sql = """SELECT `id` FROM `word_files` WHERE `word_name` = %s AND `user`= %s """
        select_re = cur.execute(select_sql, (word_name, user))
        cur.close()
        conn.close()
        if select_re == 0:
            app.logger.error("该用户没有该文件")
            return 3
        elif select_re != 0:
            delete_sql = """DELETE FROM `word_files` WHERE `word_name` = %s """
            conn = getconn()
            cur = conn.cursor()
            re = cur.execute(delete_sql, (word_name))
            if re < 1:
                conn.rollback()
                cur.close()
                conn.close()
                return 0
            else:
                try:
                    os.remove(os.path.join(upload_files, file_name))
                except Exception as e:
                    return 0
                return 1
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


# 查询库中该用户能看到的所有文件
def select_file_all(user_name):
    conn = getconn()
    cur = conn.cursor()
    try:
        user_name = user_name
        is_private = 0
        select_sql = """SELECT `id`, `user`, `word_name`, `file_label`, `is_private`  FROM `word_files` WHERE `user` = %s OR `is_private`= %s """
        re = cur.execute(select_sql, (user_name, is_private))
        if re < 1:
            return 0
        result = cur.fetchall()
        file_list = []
        for item in result:
            file_dict = {}
            file_dict['user_name'] = item[1]
            file_dict['word_name'] = item[2]
            # if item[3]:
            #     try:
            #         result_info = item[3].split(',')
            #         j = 0
            #         for i in result_info:
            #             j = j + 1
            #             file_dict['file_label' + str(j)] = i
            #     except Exception as e:
            #         file_dict['file_label'] = item[3]
            # else:
            #     file_dict['file_label'] = item[3]
            file_dict['file_label'] = item[3].split(',')
            file_dict['is_private'] = item[4]
            file_list.append(file_dict)
        return file_list
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


# 插入message信息
def insert_message_info(user_name, model_info, message):
    conn = getconn()
    cur = conn.cursor()
    current_time = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(time.time()))
    insert_sql = """INSERT INTO `messages` (`user`,`model_info`,`message`,`time`) VALUES (%s,%s,%s,%s)"""
    try:
        insert_re = cur.execute(insert_sql, (user_name, model_info, message, current_time))
        if insert_re < 1:
            return 0
        else:
            return 1
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


# 查找所有留言消息
def select_messages():
    conn = getconn()
    cur = conn.cursor()

    try:
        select_sql = """SELECT `id`,`user` ,`model_info` ,`message` ,`time`  FROM `messages` """
        re = cur.execute(select_sql)
        if re < 1:
            return 0
        else:
            count = re
            result = cur.fetchall()
            lists = []
            for item in result:
                dict = {}
                dict['id'] = item[0]
                dict['user'] = item[1]
                dict['model_info'] = item[2]
                dict['message'] = item[3]
                dict['time'] = item[4]
                lists.append(dict)
            return lists, count

    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


def select_message_info_page(page, count):
    conn = getconn()
    cur = conn.cursor()
    try:
        select_info = '''SELECT `id`,`user` ,`model_info` ,`message` ,`time` FROM `messages` LIMIT %s, %s;'''
        count = count
        if int(page) == 1 or int(page) == 0:
            page_t = 0
        else:
            page_t = (int(page) - 1) * int(count)
        re = cur.execute(select_info, (int(page_t), int(count)))
        if re < 1:
            return 0
        else:
            result = cur.fetchall()
            lists = []
            for item in result:
                # pro_name = select_pro_by_id(item[1])
                dict = {}
                dict['id'] = item[0]
                dict['user'] = item[1]
                # dict['administrative_name'] = pro_name
                dict['model_info'] = item[2]
                dict['message'] = item[3]
                dict['time'] = item[4]
                lists.append(dict)
            return lists
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)

    pass

# 删除指定的消息
def delete_messages(user_name, message_id):
    conn = getconn()
    cur = conn.cursor()

    try:
        # 首先验证要删除的消息的用户是否是创建消息的用户
        select_sql = """SELECT `user` FROM `messages` WHERE `id`= %s"""
        delete_sql = """DELETE FROM `messages` WHERE id = %s """
        re = cur.execute(select_sql, message_id)
        if re < 1:
            return 0
        else:
            result = cur.fetchone()
            original_name = result[0]
            if original_name != user_name:
                return "您没有权限删除该消息"
            else:
                re = cur.execute(delete_sql, message_id)
                if re < 1:
                    return 0
                return 1

    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


# 登录界面
def login(name, password):
    conn = getconn()
    cur = conn.cursor()
    try:
        password = md5(password)
        sql = "select `user`, `position` from users where user=%s and passwd=%s"
        result = cur.execute(sql, (name, password))
        print(result)
        if result < 1:

            return 0, 0
        else:
            info = cur.fetchone()
            return result, info[1]

    except (Exception) as e:
        app.logger.error(e)
        return "0"
    finally:
        closeAll(conn,cur)


# 插入用户
def add_new_user(new_user,password, user_level):
    conn = getconn()
    cur = conn.cursor()
    print('add_new_user')
    try:
        insert_sql = """INSERT INTO `users` (`user`,`passwd`,`position`) VALUES (%s,%s,%s)"""
        select_sql = """SELECT `user` FROM `users` WHERE `user`= %s"""
        select_result = cur.execute(select_sql, new_user)
        if select_result != 0:
            return "当前用户已存在"
        password = md5(password)
        print(password)
        result = cur.execute(insert_sql, (new_user, password, user_level))
        if result < 1:
            return "插入失败"
        else:
            return "插入成功"
    except (Exception) as e:
        app.logger.error(e)
        return "0"
    finally:
        closeAll(conn, cur)


# 删除用户
def delete_user_info(data):
    conn = getconn()
    cur = conn.cursor()
    id = data.get('id')
    try:
        delete_sql = """DELETE FROM `users` WHERE id = %s """
        re = cur.execute(delete_sql, id)
        if re < 1:
            return 0
        else:
            return 1
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


# 查单个用户的密码
def select_user_secret(id):
    conn = getconn()
    cur = conn.cursor()
    try:
        select_sql = """SELECT `passwd` FROM `users` WHERE `id` = %s """
        re = cur.execute(select_sql, id)
        if re < 1:
            return 0
        else:
            item = cur.fetchone()
            # print(item)
            pass_word = item[0]
            return pass_word
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


# 修改用户信息
def update_users_info(datas):
    where_dict = datas.get('where_dict')
    up_dict = datas.get('up_dict')
    if len(where_dict) == 0 or len(up_dict) == 0:
        if len(where_dict) == 0:
            return "更新条件不能为空"
        if len(up_dict) == 0:
            return "更新内容不能为空"
    else:
        conn = getconn()
        cur = conn.cursor()
        try:
            where_sql = ''
            up_sql = ''
            where_num = 0
            up_num = 0
            for k, v in where_dict.items():
                ency_password = select_user_secret(v)
                print(ency_password)
                where_num += 1
                if where_num < len(where_dict):
                    where_sql += "{} = '{}' AND ".format(k, v)
                else:
                    where_sql += "{} = '{}'".format(k, v)
            for k, v in up_dict.items():
                if k == "passwd":
                    if v:
                        v = md5(v)
                    else:
                        v = ency_password
                up_num += 1
                if up_num < len(up_dict):
                    up_sql += "{} = '{}', ".format(k, v)
                else:
                    up_sql += "{} = '{}' ".format(k, v)
            update_sql = """UPDATE `users` SET """ + up_sql + ' WHERE ' + where_sql
            re = cur.execute(update_sql)
            if re < 1:
                return 0
            else:
                return 1
        except Exception as erro:
            app.logger.error(erro)
            return 0
        finally:
            closeAll(conn, cur)


# 查询所有用户信息
def select_users_all():
    conn = getconn()
    cur = conn.cursor()

    try:
        select_sql = """SELECT `id`,`user` ,`position` FROM `users` """
        re = cur.execute(select_sql)
        if re < 1:
            return 0, 0
        else:
            count = re
            result = cur.fetchall()
            lists = []
            for item in result:
                dict = {}
                dict['id'] = item[0]
                dict['user'] = item[1]
                dict['position'] = item[2]
                lists.append(dict)
            return lists, count
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn,cur)


# 获取一条用户数据
def select_user_info_one(data):
    conn = getconn()
    cur = conn.cursor()

    id = data['id']
    try:
        select_sql = """SELECT `user`,`position`  FROM `users` WHERE `id` = %s """
        re = cur.execute(select_sql, id)
        if re < 1:
            return 0
        else:
            item = cur.fetchone()
            dict = {}
            dict['user_name'] = item[0]
            dict['position'] = item[1]
            return dict
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


# 分页获取用户数据
def select_user_info_page(page, count):
    conn = getconn()
    cur = conn.cursor()
    try:
        select_info = '''SELECT `id`,`user`,`position` FROM `users` LIMIT %s, %s;'''
        count = count
        if int(page) == 1 or int(page) == 0:
            page_t = 0
        else:
            page_t = (int(page) - 1) * int(count)
        re = cur.execute(select_info, (int(page_t), int(count)))
        if re < 1:
            return 0
        else:
            result = cur.fetchall()
            lists = []
            for item in result:
                dict = {}
                dict['id'] = item[0]
                dict['user'] = item[1]
                dict['position'] = item[2]
                lists.append(dict)
            return lists
    except Exception as erro:
        app.logger.error(erro)
        return 0
    finally:
        closeAll(conn, cur)


def update_file_label(datas):
    where_dict = datas.get('where_dict')
    up_dict = datas.get('up_dict')
    print(where_dict)
    print(up_dict)
    if len(where_dict) == 0 or len(up_dict) == 0:
        if len(where_dict) == 0:
            return "更新条件不能为空"
        if len(up_dict) == 0:
            return "更新内容不能为空"
    else:
        conn = getconn()
        cur = conn.cursor()
        try:
            where_sql = ''
            up_sql = ''
            where_num = 0
            up_num = 0
            for k, v in where_dict.items():
                where_num += 1
                if where_num < len(where_dict):
                    where_sql += "{} = '{}' AND ".format(k, v)
                else:
                    where_sql += "{} = '{}'".format(k, v)
            for k, v in up_dict.items():
                up_num += 1
                if up_num < len(up_dict):
                    up_sql += "{} = '{}', ".format(k, v)
                else:
                    up_sql += "{} = '{}' ".format(k, v)
            update_sql = """UPDATE `word_files` SET """ + up_sql + ' WHERE ' + where_sql
            re = cur.execute(update_sql)
            if re < 1:
                return 0
            else:
                return 1
        except Exception as erro:
            app.logger.error(erro)
            return 0
        finally:
            closeAll(conn, cur)


